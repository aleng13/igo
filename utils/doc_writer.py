# utils/doc_writer.py

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def save_mom_to_docx(mom_text: str, output_path: str = "IEDC_MOM.docx"):
    doc = Document()

    # Title
    title = doc.add_paragraph("MINUTES OF MEETING")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    # Timestamp (just for file clarity)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")

    # Add the actual Gemini output (MoM)
    for line in mom_text.split("\n"):
        doc.add_paragraph(line)

    # Save
    doc.save(output_path)
    print(f"[✅] MoM saved to {output_path}")
